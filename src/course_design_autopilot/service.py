import os
import re
import zipfile
from pathlib import Path

from .models import CourseRequirementsResult, WorkspaceInspectionResult

_SUPPORTED_SUFFIXES = {".txt", ".md", ".pdf", ".docx", ".doc"}
_TEMPLATE_PATTERN = re.compile(r"(模板|template|格式)", re.IGNORECASE)
_FRONT_MATTER_PATTERN = re.compile(
    r"^---\nlanguage:\s*(?P<language>[^\n]+)\ntemplate:\s*(?P<template>[^\n]+)\n---",
    re.MULTILINE,
)
_CJK_PATTERN = re.compile(r"[\u4e00-\u9fff]")


def _parse_with_mineru(file_path: str) -> str:
    """Convert a PDF or DOCX file to Markdown using MinerU when available."""
    if file_path.endswith((".docx", ".doc")):
        try:
            from magic_pdf.data.data_reader_writer import FileBasedDataReader  # noqa: F401
        except ImportError:
            from docx import Document

            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)

        try:
            from magic_pdf.tools.office_pipeline import office_convert

            return office_convert(file_path)
        except ImportError:
            from docx import Document

            doc = Document(file_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)

    from magic_pdf.data.data_reader_writer import FileBasedDataReader
    from magic_pdf.data.dataset import PymuDocDataset
    from magic_pdf.model.doc_analyze_by_custom_model import doc_analyze

    reader = FileBasedDataReader("")
    pdf_bytes = reader.read(file_path)
    dataset = PymuDocDataset(pdf_bytes)
    infer_result = dataset.apply(doc_analyze, ocr=False)
    return infer_result.pipe_txt_mode(None, None).get_markdown("")


def _parse_pdf_without_mineru(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join(page for page in pages if page)
    return text or "[PDF extracted but no text content detected]"


def _detect_language(text: str) -> str:
    if not text.strip():
        return "zh"

    if _CJK_PATTERN.search(text):
        return "zh"

    ascii_letters = sum(character.isascii() and character.isalpha() for character in text)
    cjk_letters = len(_CJK_PATTERN.findall(text))
    if ascii_letters > 0 and cjk_letters == 0:
        return "en"

    try:
        from langdetect import detect

        language = detect(text[:2000])
        return "zh" if language.startswith("zh") else language
    except Exception:
        return "en"


def _find_template(workspace_path: Path) -> Path | None:
    for root, _, files in os.walk(workspace_path):
        for file_name in files:
            candidate = Path(root) / file_name
            if candidate.suffix.lower() == ".docx" and _TEMPLATE_PATTERN.search(file_name):
                return candidate.relative_to(workspace_path)
    return None


def _extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="replace")
    if suffix in {".docx", ".doc"}:
        return _parse_with_mineru(str(file_path))
    if suffix == ".pdf":
        try:
            return _parse_with_mineru(str(file_path))
        except ModuleNotFoundError as error:
            if error.name != "magic_pdf":
                raise
            return _parse_pdf_without_mineru(str(file_path))
    return f"[Unsupported file type: {file_path.name}]"


def _render_requirements(language: str, template_path: Path | None, sections: list[str]) -> str:
    template_value = "null" if template_path is None else template_path.as_posix()
    front_matter = f"---\nlanguage: {language}\ntemplate: {template_value}\n---\n\n"
    body = "# Course Requirements\n"
    if sections:
        body += "".join(sections)
    return front_matter + body


class CourseDesignService:
    def extract_course_requirements(
        self,
        archive_path: Path | str,
        workspace_path: Path | str,
    ) -> CourseRequirementsResult:
        archive = Path(archive_path)
        workspace = Path(workspace_path)

        if not archive.exists():
            raise FileNotFoundError(f"Archive not found: {archive}")

        workspace.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(archive, "r") as archive_file:
            archive_file.extractall(workspace)

        sections: list[str] = []
        source_files: list[str] = []
        for file_path in sorted(workspace.rglob("*")):
            if not file_path.is_file():
                continue
            if file_path.name == "COURSE_REQUIREMENTS.md":
                continue
            if file_path.suffix.lower() not in _SUPPORTED_SUFFIXES:
                continue

            relative_name = file_path.relative_to(workspace).as_posix()
            source_files.append(relative_name)
            sections.append(f"\n\n## {relative_name}\n\n{_extract_text(file_path)}")

        merged_text = "\n".join(sections)
        language = _detect_language(merged_text)
        template_path = _find_template(workspace)
        requirements_path = workspace / "COURSE_REQUIREMENTS.md"
        requirements_path.write_text(
            _render_requirements(language, template_path, sections),
            encoding="utf-8",
        )

        return CourseRequirementsResult(
            workspace_path=workspace,
            requirements_path=requirements_path,
            language=language,
            template_path=template_path,
            source_files=source_files,
        )

    def inspect_workspace(self, workspace_path: Path | str) -> WorkspaceInspectionResult:
        workspace = Path(workspace_path)
        requirements_path = workspace / "COURSE_REQUIREMENTS.md"
        if not requirements_path.exists():
            return WorkspaceInspectionResult(
                workspace_path=workspace,
                requirements_exists=False,
                requirements_path=requirements_path,
                language=None,
                template_path=None,
            )

        content = requirements_path.read_text(encoding="utf-8", errors="replace")
        match = _FRONT_MATTER_PATTERN.search(content)
        language = None
        template_path = None
        if match:
            language = match.group("language").strip()
            template_value = match.group("template").strip()
            if template_value != "null":
                template_path = Path(template_value)

        return WorkspaceInspectionResult(
            workspace_path=workspace,
            requirements_exists=True,
            requirements_path=requirements_path,
            language=language,
            template_path=template_path,
        )
